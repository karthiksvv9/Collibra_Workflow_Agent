from __future__ import annotations

import json
import re
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET


SUPPORTED_EXTENSIONS = {".docx", ".pdf", ".xlsx", ".xml", ".bpmn", ".app", ".form", ".json", ".txt", ".md", ".zip"}


@dataclass(slots=True)
class SourceDocument:
    path: str
    kind: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


def discover_documents(root: Path) -> list[Path]:
    if not root.exists():
        return []
    return sorted(path for path in root.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS)


def load_document(path: str | Path) -> SourceDocument:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return _load_docx(file_path)
    if suffix == ".pdf":
        return _load_pdf(file_path)
    if suffix == ".xlsx":
        return _load_xlsx(file_path)
    if suffix in {".xml", ".bpmn"}:
        return _load_xml(file_path, kind=suffix[1:])
    if suffix in {".app", ".form", ".json"}:
        return _load_json_or_text(file_path, kind=suffix[1:])
    if suffix == ".zip":
        return _load_workflow_zip(file_path)
    return SourceDocument(str(file_path), suffix[1:] or "text", file_path.read_text(encoding="utf-8", errors="ignore"))


def _load_docx(path: Path) -> SourceDocument:
    try:
        from docx import Document as DocxDocument

        document = DocxDocument(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return SourceDocument(str(path), "docx", text, {"paragraphs": len(document.paragraphs)})
    except Exception:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml")
        root = ET.fromstring(xml)
        paragraphs = ["".join(node.itertext()).strip() for node in root.iter() if node.tag.endswith("}p")]
        return SourceDocument(str(path), "docx", "\n".join(filter(None, paragraphs)), {"paragraphs": len(paragraphs)})


def _load_pdf(path: Path) -> SourceDocument:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return SourceDocument(str(path), "pdf", "\n".join(pages), {"pages": len(pages)})
    except Exception as exc:
        return SourceDocument(str(path), "pdf", "", {"warning": f"PDF extraction failed: {exc}"})


def _load_xlsx(path: Path) -> SourceDocument:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(path, data_only=True, read_only=True)
        parts: list[str] = []
        sheets: dict[str, dict[str, Any]] = {}
        for sheet in workbook.worksheets:
            rows = list(sheet.iter_rows(values_only=True))
            sheets[sheet.title] = {"rows": len(rows), "columns": sheet.max_column}
            if not rows:
                continue
            header = [str(value).strip() if value is not None else "" for value in rows[0]]
            parts.append(f"# Sheet: {sheet.title}\nColumns: {', '.join(filter(None, header))}")
            for row in rows[1:51]:
                values = ["" if value is None else str(value) for value in row]
                parts.append(" | ".join(values))
        return SourceDocument(str(path), "xlsx", "\n".join(parts), {"sheets": sheets})
    except Exception:
        return _load_xlsx_zip(path)


def _load_xlsx_zip(path: Path) -> SourceDocument:
    with zipfile.ZipFile(path) as archive:
        shared_strings = _read_shared_strings(archive)
        workbook = archive.read("xl/workbook.xml")
        workbook_root = ET.fromstring(workbook)
        rels = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        relation_targets = {
            rel.attrib["Id"]: rel.attrib["Target"]
            for rel in rels
            if rel.attrib.get("Target", "").startswith("worksheets/")
        }
        namespaces = {
            "main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
            "rel": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        }
        parts: list[str] = []
        sheets: dict[str, dict[str, Any]] = {}
        for sheet in workbook_root.findall(".//main:sheet", namespaces):
            name = sheet.attrib.get("name", "Sheet")
            rel_id = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            target = relation_targets.get(rel_id or "")
            if not target:
                continue
            rows = _read_sheet_rows(archive, f"xl/{target}", shared_strings)
            sheets[name] = {"rows": len(rows), "columns": max((len(row) for row in rows), default=0)}
            parts.append(f"# Sheet: {name}")
            for row in rows[:51]:
                parts.append(" | ".join(row))
        return SourceDocument(str(path), "xlsx", "\n".join(parts), {"sheets": sheets})


def _read_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
    return ["".join(item.itertext()) for item in root]


def _read_sheet_rows(archive: zipfile.ZipFile, name: str, shared_strings: list[str]) -> list[list[str]]:
    root = ET.fromstring(archive.read(name))
    rows: list[list[str]] = []
    for row in root.iter():
        if not row.tag.endswith("}row"):
            continue
        values: list[str] = []
        for cell in row:
            if not cell.tag.endswith("}c"):
                continue
            value_node = next((child for child in cell if child.tag.endswith("}v")), None)
            raw = value_node.text if value_node is not None else ""
            if cell.attrib.get("t") == "s" and raw.isdigit():
                values.append(shared_strings[int(raw)])
            else:
                values.append(raw or "")
        rows.append(values)
    return rows


def _load_xml(path: Path, kind: str) -> SourceDocument:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    metadata: dict[str, Any] = {}
    try:
        root = ET.fromstring(raw)
        metadata = {
            "root": _strip_ns(root.tag),
            "elements": len(list(root.iter())),
            "ids": sorted(set(filter(None, (node.attrib.get("id") for node in root.iter()))))[:250],
        }
        text = _xml_to_semantic_text(root)
    except Exception as exc:
        metadata = {"warning": f"XML parse failed: {exc}"}
        text = raw
    return SourceDocument(str(path), kind, text, metadata)


def _load_json_or_text(path: Path, kind: str) -> SourceDocument:
    raw = path.read_text(encoding="utf-8", errors="ignore")
    try:
        data = json.loads(raw)
        return SourceDocument(str(path), kind, json.dumps(data, indent=2, sort_keys=True), {"json_type": type(data).__name__})
    except json.JSONDecodeError:
        return SourceDocument(str(path), kind, raw)


def _load_workflow_zip(path: Path) -> SourceDocument:
    parts = [f"# Workflow ZIP: {path.name}"]
    metadata: dict[str, Any] = {"members": 0, "bpmn": 0, "forms": 0, "apps": 0, "groovy": 0}
    with zipfile.ZipFile(path) as archive:
        names = [name for name in archive.namelist() if not name.endswith("/")]
        metadata["members"] = len(names)
        for name in names:
            suffix = Path(name).suffix.lower()
            if suffix not in {".bpmn", ".xml", ".form", ".app", ".json", ".groovy", ".md", ".txt"}:
                continue
            raw = archive.read(name).decode("utf-8", errors="ignore")
            if suffix in {".bpmn", ".xml"}:
                metadata["bpmn"] += 1
                parts.append(f"\n## BPMN/XML: {name}\n{_summarize_xml_text(raw)}")
            elif suffix == ".form":
                metadata["forms"] += 1
                parts.append(f"\n## Form: {name}\n{_summarize_json_text(raw)}")
            elif suffix == ".app":
                metadata["apps"] += 1
                parts.append(f"\n## App: {name}\n{_summarize_json_text(raw)}")
            elif suffix == ".groovy":
                metadata["groovy"] += 1
                parts.append(f"\n## Groovy: {name}\n{raw[:4000]}")
            else:
                parts.append(f"\n## {name}\n{raw[:4000]}")
    return SourceDocument(str(path), "zip", "\n".join(parts), metadata)


def _summarize_json_text(raw: str) -> str:
    try:
        data = json.loads(raw)
        return json.dumps(data, indent=2, sort_keys=True)[:5000]
    except json.JSONDecodeError:
        return raw[:5000]


def _summarize_xml_text(raw: str) -> str:
    try:
        root = ET.fromstring(raw)
        return _xml_to_semantic_text(root)[:6000]
    except Exception:
        return raw[:6000]


def _xml_to_semantic_text(root: ET.Element) -> str:
    lines: list[str] = []
    for node in root.iter():
        tag = _strip_ns(node.tag)
        label = node.attrib.get("name") or node.attrib.get("id") or tag
        attrs = " ".join(f"{_strip_ns(key)}={value}" for key, value in sorted(node.attrib.items()))
        text = " ".join(node.itertext()).strip()
        text = re.sub(r"\s+", " ", text)
        lines.append(f"{tag}: {label} {attrs} {text[:500]}".strip())
    return "\n".join(lines)


def _strip_ns(value: str) -> str:
    return value.rsplit("}", 1)[-1] if "}" in value else value
